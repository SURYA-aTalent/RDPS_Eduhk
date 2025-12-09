#!/usr/bin/env python3
"""Convert markdown documentation to Word document with formatting."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_line(line):
    """Parse a markdown line and return its type and content."""
    # Headers
    if line.startswith('# '):
        return ('h1', line[2:].strip())
    elif line.startswith('## '):
        return ('h2', line[3:].strip())
    elif line.startswith('### '):
        return ('h3', line[4:].strip())
    elif line.startswith('#### '):
        return ('h4', line[5:].strip())
    # Bold
    elif line.startswith('**') and line.endswith('**'):
        return ('bold', line[2:-2].strip())
    # Horizontal rule
    elif line.strip() in ['---', '***', '___']:
        return ('hr', '')
    # Code block
    elif line.startswith('```'):
        return ('code_fence', line[3:].strip())
    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        return ('bullet', line[2:].strip())
    # Numbered list
    elif re.match(r'^\d+\.\s', line):
        return ('numbered', re.sub(r'^\d+\.\s', '', line).strip())
    # Table row
    elif line.startswith('|') and line.endswith('|'):
        return ('table', line)
    # Empty line
    elif not line.strip():
        return ('empty', '')
    # Regular paragraph
    else:
        return ('paragraph', line.strip())

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, code, etc.)."""
    # Handle bold text **text**
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def create_docx_from_markdown(md_file, docx_file):
    """Convert markdown file to formatted Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        line_type, content = parse_markdown_line(line)

        # Handle code blocks
        if line_type == 'code_fence':
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End code block
                in_code_block = False
                p = doc.add_paragraph()
                for code_line in code_lines:
                    run = p.add_run(code_line + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        elif in_code_block:
            code_lines.append(line)

        # Handle tables
        elif line_type == 'table':
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        elif in_table and line_type == 'empty':
            # End table
            in_table = False
            # Parse and create table
            rows = [row for row in table_lines if not all(c in '|-: ' for c in row)]
            if rows:
                cells = [[cell.strip() for cell in row.split('|')[1:-1]] for row in rows]
                if cells:
                    table = doc.add_table(rows=len(cells), cols=len(cells[0]))
                    table.style = 'Light Grid Accent 1'
                    for row_idx, row_data in enumerate(cells):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
                            if row_idx == 0:
                                # Header row
                                table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
            table_lines = []

        # Handle other elements
        elif line_type == 'h1':
            p = doc.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_type == 'h2':
            doc.add_heading(content, level=2)
        elif line_type == 'h3':
            doc.add_heading(content, level=3)
        elif line_type == 'h4':
            doc.add_heading(content, level=4)
        elif line_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
        elif line_type == 'hr':
            doc.add_paragraph('_' * 80)
        elif line_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
        elif line_type == 'numbered':
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
        elif line_type == 'paragraph' and content:
            p = doc.add_paragraph()
            add_formatted_text(p, content)
        elif line_type == 'empty':
            pass  # Skip empty lines

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Document created: {docx_file}")

if __name__ == '__main__':
    md_file = '/Users/praveen/Documents/projects/eduhk_rdps/Production_Database_Import_Documentation.md'
    docx_file = '/Users/praveen/Documents/projects/eduhk_rdps/Production_Database_Import_Documentation.docx'
    create_docx_from_markdown(md_file, docx_file)
